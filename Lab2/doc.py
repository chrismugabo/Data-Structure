from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading('Lab 2: Asymptotic Complexity of an Algorithm', level=1)

# Add the task descriptions and answers
doc.add_heading('Task 1:', level=2)
doc.add_paragraph(
    "Answer the following questions about the time complexity of given expressions "
    "and visualize their growth rates."
)

# Part a
doc.add_heading('a. Is (n+5)^2 = O(n log n)?', level=3)
doc.add_paragraph(
    "(n+5)^2 simplifies to n^2 + 10n + 25, which is dominated by n^2. Since quadratic growth (n^2) "
    "is much faster than linearithmic growth (n log n), (n+5)^2 cannot be considered O(n log n). "
    "Quadratic growth outpaces linearithmic growth as n increases."
)

# Part b
doc.add_heading('b. Is n^(14/16) = O(n log n)?', level=3)
doc.add_paragraph(
    "n^(14/16) or n^0.875 still represents polynomial growth, albeit less than linear growth (n^1). "
    "Polynomial growth tends to surpass n log n because it increases faster with n, making n^(14/16) not O(n log n)."
)

# Part c
doc.add_heading('c. What is the time complexity of n+5, (n+5)^3, and (n+5)^5?', level=3)
doc.add_paragraph(
    "n+5: This expression is linear, so its complexity is O(n). "
    "(n+5)^3 and (n+5)^5: These are polynomial expressions where the highest powers are n^3 and n^5 respectively, "
    "leading to complexities of O(n^3) and O(n^5)."
)

# Save the document
doc_path = '/mnt/data/Asymptotic_Complexity_Lab_Report.docx'
doc.save(doc_path)
doc_path
